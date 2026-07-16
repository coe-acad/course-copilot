"""
Utility helpers to convert rich text (markdown) into DOCX documents.
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import List, Optional, Union
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as _RT

from app.utils.markdown_media import fetch_image_stream, parse_image_line
from app.utils.latex_to_text import latex_to_text

DocxBytes = bytes
PathLike = Union[str, Path]


def text_to_docx(
    text: str,
    *,
    output_path: Optional[PathLike] = None,
) -> Union[DocxBytes, Path]:
    """Convert plain/markdown text into a professionally formatted DOCX document."""

    if not text:
        raise ValueError("text must be a non-empty string")

    # Convert LaTeX math to readable plain text so equations don't render raw.
    text = latex_to_text(text)

    doc = Document()

    # Set default page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    _apply_default_styles(doc)
    _ensure_hyperlink_style(doc)
    _markdown_to_docx(doc, text)

    buffer = BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    if output_path:
        path_obj = Path(output_path)
        path_obj.write_bytes(docx_bytes)
        return path_obj

    return docx_bytes


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def _add_abstract_num(numbering, abstract_id: int, num_fmt: str, lvl_text: str, font: Optional[str]):
    """Append a single-level ``<w:abstractNum>`` definition to numbering.xml."""
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    text_el = OxmlElement("w:lvlText")
    text_el.set(qn("w:val"), lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    for child in (start, fmt, text_el, jc, ppr):
        lvl.append(child)
    if font:
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), font)
        rfonts.set(qn("w:hAnsi"), font)
        rfonts.set(qn("w:hint"), "default")
        rpr.append(rfonts)
        lvl.append(rpr)
    abstract.append(lvl)
    # All <w:abstractNum> must precede every <w:num> in the schema, so insert
    # abstracts at the front (nums are always appended at the end).
    numbering.insert(0, abstract)


def _add_num(numbering, num_id: int, abstract_id: int):
    """Append a ``<w:num>`` pointing at an abstract definition."""
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def _init_list_numbering(doc: Document) -> dict:
    """Create the shared bullet + decimal abstract definitions once per document.

    Applying the "List Bullet"/"List Number" *style* alone does not attach a
    numbering definition, so Word shows no bullet/number glyph. We add real
    ``<w:abstractNum>`` definitions and cache state (the shared bullet num id,
    the decimal abstract id, and the next free num id) on the document.
    """
    state = getattr(doc, "_list_num_state", None)
    if state is not None:
        return state

    numbering = doc.part.numbering_part.element  # <w:numbering>

    # Pick IDs that don't collide with anything already in the template.
    existing_abstract = [
        int(e.get(qn("w:abstractNumId")))
        for e in numbering.findall(qn("w:abstractNum"))
        if e.get(qn("w:abstractNumId")) is not None
    ]
    existing_num = [
        int(e.get(qn("w:numId")))
        for e in numbering.findall(qn("w:num"))
        if e.get(qn("w:numId")) is not None
    ]
    bullet_abstract = max(existing_abstract, default=0) + 1
    decimal_abstract = bullet_abstract + 1
    _add_abstract_num(numbering, bullet_abstract, "bullet", "•", "Symbol")
    _add_abstract_num(numbering, decimal_abstract, "decimal", "%1.", None)

    # Bullets don't count, so all bullet lists can share one num. Numbered lists
    # each get a fresh num (see _new_ordered_num_id) so their counters restart.
    bullet_num_id = max(existing_num, default=0) + 1
    _add_num(numbering, bullet_num_id, bullet_abstract)

    state = {
        "numbering": numbering,
        "decimal_abstract": decimal_abstract,
        "bullet_num_id": bullet_num_id,
        "next_num_id": bullet_num_id + 1,
    }
    doc._list_num_state = state
    return state


def _bullet_num_id(doc: Document) -> int:
    """Return the shared bullet numbering id (bullets never need to restart)."""
    return _init_list_numbering(doc)["bullet_num_id"]


def _new_ordered_num_id(doc: Document) -> int:
    """Allocate a fresh num id for one ordered list so its count restarts at 1."""
    state = _init_list_numbering(doc)
    num_id = state["next_num_id"]
    state["next_num_id"] += 1
    _add_num(state["numbering"], num_id, state["decimal_abstract"])
    return num_id


def _apply_list_number(paragraph, num_id: int):
    """Attach direct numbering (numPr) so the bullet/number actually renders."""
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(num)
    ppr.append(numpr)


def _apply_default_styles(doc: Document):
    """Configure the document's built-in styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def _set_paragraph_color(para, rgb: RGBColor):
    for run in para.runs:
        run.font.color.rgb = rgb


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------

def _markdown_to_docx(doc: Document, text: str):
    """Parse markdown blocks and add them to the document."""
    for block in _split_blocks(text):
        b_type = block.get("type")

        if b_type == "header1":
            p = doc.add_heading(block["text"], level=1)
            p.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            p.runs[0].font.size = Pt(20)
            p.runs[0].bold = True

        elif b_type == "header2":
            p = doc.add_heading(block["text"], level=2)
            p.runs[0].font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            p.runs[0].font.size = Pt(16)
            p.runs[0].bold = True

        elif b_type == "header3":
            p = doc.add_heading(block["text"], level=3)
            p.runs[0].font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            p.runs[0].font.size = Pt(13)
            p.runs[0].bold = True

        elif b_type == "ul":
            bullet_num_id = _bullet_num_id(doc)
            for item in block.get("items", []):
                p = doc.add_paragraph(style="List Bullet")
                _apply_list_number(p, bullet_num_id)
                _add_inline_runs(p, item)

        elif b_type == "ol":
            number_num_id = _new_ordered_num_id(doc)  # fresh id -> restarts at 1
            for item in block.get("items", []):
                p = doc.add_paragraph(style="List Number")
                _apply_list_number(p, number_num_id)
                _add_inline_runs(p, item)

        elif b_type == "code":
            p = doc.add_paragraph()
            run = p.add_run(block["text"])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            # Light grey shading
            _shade_paragraph(p, "F3F4F6")

        elif b_type == "quote":
            quote_text = " ".join(block.get("lines", []))
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(quote_text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        elif b_type == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            col_count = max(len(headers), max((len(r) for r in rows), default=0))
            if col_count == 0:
                continue
            table = doc.add_table(rows=1 + len(rows), cols=col_count)
            table.style = "Table Grid"

            # Header row
            hdr_row = table.rows[0]
            for ci, cell_text in enumerate(headers):
                cell = hdr_row.cells[ci]
                p = cell.paragraphs[0]
                _add_inline_runs(p, cell_text.strip())
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
                _shade_cell(cell, "F9FAFB")

            # Data rows
            for ri, row_cells in enumerate(rows):
                tbl_row = table.rows[ri + 1]
                for ci in range(col_count):
                    cell_text = row_cells[ci].strip() if ci < len(row_cells) else ""
                    p = tbl_row.cells[ci].paragraphs[0]
                    _add_inline_runs(p, cell_text)

        elif b_type == "image":
            _add_image(doc, block)

        elif b_type == "rule":
            _add_horizontal_rule(doc)

        else:
            # Regular paragraph
            lines = block.get("lines", [])
            paragraph_text = " ".join(line.strip() for line in lines if line.strip())
            if paragraph_text:
                p = doc.add_paragraph()
                _add_inline_runs(p, paragraph_text)


# ---------------------------------------------------------------------------
# Inline markdown → runs
# ---------------------------------------------------------------------------

# Bare URL autolink: http(s):// or www. up to the next space/bracket/quote.
_URL_RE = re.compile(r"(?:https?://|www\.)[^\s<>()\[\]\"']+", re.IGNORECASE)


def _split_trailing_punct(url: str):
    """Split sentence punctuation that a bare URL match greedily swallowed.

    ``https://x.com/page.`` -> (``https://x.com/page``, ``.``). Also releases a
    dangling ``)`` when the URL has no matching ``(``.
    """
    trail = ""
    while url and url[-1] in ".,!?":
        trail = url[-1] + trail
        url = url[:-1]
    while url.endswith(")") and url.count("(") < url.count(")"):
        trail = ")" + trail
        url = url[:-1]
    return url, trail


def _ensure_hyperlink_style(doc: Document):
    """Create a blue+underlined 'Hyperlink' character style if the template lacks one."""
    from docx.enum.style import WD_STYLE_TYPE

    if any(s.name == "Hyperlink" for s in doc.styles):
        return
    style = doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    style.font.underline = True


def _add_hyperlink(para, text: str, url: str):
    """Add a clickable external hyperlink run (blue, underlined) to a paragraph."""
    part = para.part
    r_id = part.relate_to(url, _RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    # Reference the Hyperlink character style AND set explicit color/underline so
    # the link looks right whether or not the viewer honours the style.
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(rstyle)
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)

    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)

    hyperlink.append(run)
    para._p.append(hyperlink)


_INLINE_PATTERN = re.compile(
    r'(?P<link>\[(?P<ltext>[^\]]+)\]\((?P<lurl>[^)\s]+)\))'  # [text](url)
    r'|(?P<url>(?:https?://|www\.)[^\s<>\[\]"\']+)'            # bare url autolink
    r'|(?P<bi>\*\*\*(?P<bitext>.+?)\*\*\*)'                    # bold + italic
    r'|(?P<b>\*\*(?P<btext>.+?)\*\*)'                          # bold
    r'|(?P<i>\*(?P<itext>.+?)\*)'                              # italic
    r'|(?P<code>`(?P<ctext>[^`]+)`)'                           # inline code
    r'|(?P<strike>~~(?P<stext>.+?)~~)'                         # strikethrough
)


def _add_inline_runs(para, text: str):
    """Parse inline markdown and add styled runs to a paragraph."""
    pos = 0
    for m in _INLINE_PATTERN.finditer(text):
        # Add plain text before this match
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

        if m.group("link"):
            _add_hyperlink(para, m.group("ltext"), m.group("lurl"))
        elif m.group("url"):  # bare URL -> autolink
            url, trail = _split_trailing_punct(m.group("url"))
            href = url if url.lower().startswith("http") else "https://" + url
            _add_hyperlink(para, url, href)
            if trail:
                run = para.add_run(trail)
                run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        elif m.group("bi"):
            run = para.add_run(m.group("bitext"))
            run.bold = True
            run.italic = True
        elif m.group("b"):
            run = para.add_run(m.group("btext"))
            run.bold = True
        elif m.group("i"):
            run = para.add_run(m.group("itext"))
            run.italic = True
        elif m.group("code"):
            run = para.add_run(m.group("ctext"))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        elif m.group("strike"):
            run = para.add_run(m.group("stext"))
            # python-docx has no direct strikethrough property; set it via XML.
            run.font._element.get_or_add_rPr().append(OxmlElement('w:strike'))

        pos = m.end()

    # Remaining plain text
    if pos < len(text):
        run = para.add_run(text[pos:])
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)


# ---------------------------------------------------------------------------
# Block parser (reuses logic from text_to_pdf.py)
# ---------------------------------------------------------------------------

def _split_blocks(text: str) -> List[dict]:
    blocks: List[dict] = []
    lines = text.splitlines()
    idx = 0
    total = len(lines)

    while idx < total:
        line = lines[idx].rstrip("\n")
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            idx += 1
            code_lines = []
            while idx < total and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx].rstrip())
                idx += 1
            if idx < total:
                idx += 1
            blocks.append({"type": "code", "text": "\n".join(code_lines)})
            continue

        if stripped == "":
            idx += 1
            continue

        if stripped in {"---", "***", "___"} or re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", stripped):
            blocks.append({"type": "rule"})
            idx += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while idx < total and lines[idx].strip().startswith(">"):
                quote_lines.append(lines[idx].strip()[1:].strip())
                idx += 1
            blocks.append({"type": "quote", "lines": quote_lines})
            continue

        header_match = re.match(r"^(#{3,6})\s+(.*)", stripped)
        if header_match:
            blocks.append({"type": "header3", "text": header_match.group(2).strip()})
            idx += 1
            continue

        header2_match = re.match(r"^#{2}\s+(.*)", stripped)
        if header2_match:
            blocks.append({"type": "header2", "text": header2_match.group(1).strip()})
            idx += 1
            continue

        header1_match = re.match(r"^#\s+(.*)", stripped)
        if header1_match:
            blocks.append({"type": "header1", "text": header1_match.group(1).strip()})
            idx += 1
            continue

        if re.match(r"^\s*[-*+]\s+", line) and not re.match(r"^\d+\s+", stripped):
            items: List[str] = []
            while idx < total:
                current = lines[idx]
                current_stripped = current.strip()
                if current_stripped == "":
                    idx += 1
                    continue
                if not re.match(r"^\s*[-*+]\s+", current) or re.match(r"^\d+\s+", current_stripped):
                    break
                items.append(re.sub(r"^\s*[-*+]\s+", "", current).strip())
                idx += 1
            if items:
                blocks.append({"type": "ul", "items": items})
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            items = []
            while idx < total:
                current = lines[idx]
                current_stripped = current.strip()
                if current_stripped == "":
                    idx += 1
                    continue
                if not re.match(r"^\s*\d+\.\s+", current):
                    break
                items.append(re.sub(r"^\s*\d+\.\s+", "", current).strip())
                idx += 1
            if items:
                blocks.append({"type": "ol", "items": items})
            continue

        # Standalone images: ![alt](src)
        image_match = parse_image_line(stripped)
        if image_match:
            alt, src = image_match
            blocks.append({"type": "image", "alt": alt, "src": src})
            idx += 1
            continue

        # Markdown tables — lines starting with |
        if stripped.startswith("|"):
            table_lines = []
            while idx < total and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].strip())
                idx += 1

            def _parse_row(row_str):
                # Strip leading/trailing pipes then split on |
                cells = row_str.strip().strip("|").split("|")
                return [c.strip() for c in cells]

            def _is_separator(row_str):
                return all(re.match(r"^:?-+:?$", c.strip()) for c in row_str.strip().strip("|").split("|") if c.strip())

            headers = []
            rows = []
            for i, tl in enumerate(table_lines):
                if i == 0:
                    headers = _parse_row(tl)
                elif _is_separator(tl):
                    continue
                else:
                    rows.append(_parse_row(tl))

            if headers:
                blocks.append({"type": "table", "headers": headers, "rows": rows})
            continue

        paragraph_lines = [line]
        idx += 1
        while idx < total:
            lookahead = lines[idx]
            stripped_la = lookahead.strip()
            if not stripped_la:
                idx += 1
                break
            if (
                stripped_la.startswith("```")
                or stripped_la in {"---", "***", "___"}
                or (re.match(r"^\s*[-*+]\s+", lookahead) and not re.match(r"^\d+\s+", stripped_la))
                or re.match(r"^\s*\d+\.\s+", lookahead)
                or stripped_la.startswith("#")
                or stripped_la.startswith(">")
                or stripped_la.startswith("|")
                or parse_image_line(stripped_la)
            ):
                break
            paragraph_lines.append(lookahead)
            idx += 1
        blocks.append({"type": "paragraph", "lines": paragraph_lines})

    return blocks


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def _shade_paragraph(para, hex_color: str):
    """Add a background shading to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def _shade_cell(cell, hex_color: str):
    """Add a background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _add_image(doc: Document, block: dict):
    """Embed a markdown image into the document.

    Falls back to italic alt text/link if the image cannot be fetched, so a
    broken/expired image never breaks the whole document.
    """
    src = block.get("src", "")
    alt = block.get("alt", "")
    stream = fetch_image_stream(src)
    if stream is not None:
        try:
            para = doc.add_paragraph()
            para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
            shape = para.add_run().add_picture(stream)
            # Constrain to the page's content width, preserving aspect ratio.
            section = doc.sections[-1]
            content_width = section.page_width - section.left_margin - section.right_margin
            if shape.width and shape.width > content_width:
                ratio = content_width / shape.width
                shape.width = Emu(int(shape.width * ratio))
                shape.height = Emu(int(shape.height * ratio))
            return
        except Exception:
            pass

    fallback = alt or src
    p = doc.add_paragraph()
    run = p.add_run(f"[Image: {fallback}]" if fallback else "[Image]")
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def _add_horizontal_rule(doc: Document):
    """Add a horizontal rule (border bottom on an empty paragraph)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D1D5DB')
    pBdr.append(bottom)
    pPr.append(pBdr)
